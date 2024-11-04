import re
import time
import os
from langdetect import detect
from collections import defaultdict
from docx import Document
from io import BytesIO
import pandas as pd
import plotly.express as px
from  datetime import datetime
from pathlib import Path
common_german_words = set([
    "und", "der", "die", "das", "ist", "du", "ich", "nicht", "es", "sie", "wir", "ein", "er", "in", "zu", "haben",
    "auch", "ja", "nein", "mit", "wie", "auf", "für", "den", "von", "dem", "sind", "dass", "ich", "mir", "mich", "sein",
    "bei", "um", "aber", "hier", "da", "wenn", "nur", "dann", "war", "noch", "nach", "wird", "kann", "können", "wir"
    # Add more words as needed
])


def get_file_modified_time(file_path):
    # Get file metadata
    file_info = Path(file_path)

    # Get the modification time (in seconds since epoch)
    modified_time = file_info.stat().st_mtime

    # Convert to a human-readable format
    modified_time = datetime.fromtimestamp(modified_time).strftime('%Y-%m-%d %H:%M:%S')

    return modified_time
def is_gibberish(text):
    # Remove non-alphabetic characters, except for spaces and common German punctuation
    cleaned_text = re.sub(r'[^A-Za-zäöüÄÖÜß ]+', '', text)

    # Check if the detected language is German
    try:
        detected_language = detect(cleaned_text)
        if detected_language != 'de':
            return True
    except:
        return True

    # Split the text into words
    words = cleaned_text.split()

    # Check if a significant number of words are not in the common German words list
    non_german_word_count = sum(1 for word in words if word.lower() not in common_german_words)

    # Define a threshold for gibberish detection (e.g., more than half non-German words)
    threshold = 0.5
    if len(words) > 0 and (non_german_word_count / len(words)) > threshold:
        return True

    return False
def sanitize_string(input_string):
    # Remove non-printable/control characters using regex
    sanitized_string = re.sub(r'[\x00-\x1F\x7F-\x9F]', '', input_string)

    # Optionally, replace multiple newlines or spaces with a single space
    sanitized_string = re.sub(r'\s+', ' ', sanitized_string).strip()

    return sanitized_string
def sanitize_input(doc_list):
    # Remove control characters
    sanitized_docs = [re.sub(r'[\x00-\x1f\x7f-\x9f]', '', doc) for doc in doc_list]
    # Remove empty or whitespace-only docs
    sanitized_docs = [doc for doc in sanitized_docs if doc.strip()]
    # Ensure UTF-8 encoding
    sanitized_docs = [doc.encode('utf-8', 'ignore').decode('utf-8') for doc in sanitized_docs]
    return sanitized_docs
def track_files_in_directory(directory):
    file_data = []

    # Walk through directory and subdirectories
    for root, dirs, files in os.walk(directory):
        for file in files:
            file_path = os.path.join(root, file)
            file_size = os.path.getsize(file_path)  # Get file size in bytes
            file_ext = os.path.splitext(file)[1]  # Get file extension
            file_creation_time = time.ctime(os.path.getctime(file_path))  # Get file creation time

            # Append file details to list
            file_data.append({
                'File Name': file,
                'File Path': file_path,
                'File Size (Bytes)': file_size,
                'File Extension': file_ext,
                'Creation Time': file_creation_time,
                'Manipulation Time':get_file_modified_time(file_path)
            })

    # Create a DataFrame from the list
    df = pd.DataFrame(file_data)
    return df
def get_file_sizes(directory):
    file_data = []

    # Walk through the directory and collect file information
    for root, dirs, files in os.walk(directory):
        for file in files:
            file_path = os.path.join(root, file)
            file_size = os.path.getsize(file_path)  # Get file size in bytes
            relative_path = os.path.relpath(file_path, directory)  # Get relative file path

            file_data.append({
                'File Name': file,
                'File Path': relative_path,
                'File Size (Bytes)': file_size,
                'Directory': os.path.basename(root)  # Directory name
            })

    # Convert the data to a pandas DataFrame
    df = pd.DataFrame(file_data)
    return df
def create_treemap(df):
    # Create a Plotly treemap with file sizes
    fig = px.treemap(
        df,
        path=['Directory', 'File Name'],  # Hierarchical view
        values='File Size (Bytes)',  # Size of boxes based on file size
        title='File Sizes Treemap',
        color='File Size (Bytes)',  # Color based on file size
        hover_data={'File Size (Bytes)': ':.2f'}  # Show file size on hover
    )

    # Show the figure
    #fig.write_html('./Laufwerk_structure.html')


def create_word_doc(inputs, answers, documents,al_docs):
    doc = Document()

    # Add title
    doc.add_heading('Session Information', level=1)

    # Add Input Section
    doc.add_heading('Input Prompts', level=2)
    for i, item in enumerate(inputs, 1):
        doc.add_paragraph(f"{i}. {item}")

    # Add Answer Section
    doc.add_heading('Generated Answers', level=2)
    for i, item in enumerate(answers, 1):
        doc.add_paragraph(f"{i}. {item}")

    # Add Document Section
    doc.add_heading('Most Relevant Documents', level=2)
    for i, item in enumerate(documents, 1):
        doc.add_paragraph(f"{i}. {item}")

    doc.add_heading('All associated Documents', level=2)
    for i, item in enumerate(al_docs, 1):
        doc.add_paragraph(f"{i}. {str(item)}")

    # Save to BytesIO object
    byte_io = BytesIO()
    doc.save(byte_io)
    byte_io.seek(0)

    return byte_io


from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Length
from io import BytesIO
from typing import List, Dict, Any, Optional
from datetime import datetime
import os

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from io import BytesIO
from typing import List, Dict, Any, Optional
from datetime import datetime
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.oxml.shared import OxmlElement, qn
from datetime import datetime
import os


class RAGReportGenerator:
    def __init__(self, title: str = "Jahresbericht", year: str = str(datetime.now().year)):
        self.doc = Document()
        self.title = title
        self.year = year
        self.sections = {}
        self._setup_styles()
        self._setup_page_layout()

    def _create_element(self, name: str) -> OxmlElement:
        """Create a custom XML element"""
        return OxmlElement(name)

    def _setup_styles(self):
        """Setup custom styles for the document"""
        styles = self.doc.styles

        # Custom TOC style
        if 'Contents' not in styles:
            toc_style = styles.add_style('Contents', WD_STYLE_TYPE.PARAGRAPH)
            toc_style.font.name = 'Calibri'
            toc_style.font.size = Pt(11)
            toc_style.paragraph_format.left_indent = Inches(0.5)

        # Heading styles
        h1_style = styles['Heading 1']
        h1_style.font.size = Pt(16)
        h1_style.font.bold = True
        h1_style.font.color.rgb = RGBColor(0, 51, 102)

        h2_style = styles['Heading 2']
        h2_style.font.size = Pt(14)
        h2_style.font.bold = True
        h2_style.font.color.rgb = RGBColor(51, 51, 51)

        # Normal text style
        normal_style = styles['Normal']
        normal_style.font.name = 'Calibri'
        normal_style.font.size = Pt(11)

        # Highlight style
        if 'Highlight' not in styles:
            highlight_style = styles.add_style('Highlight', WD_STYLE_TYPE.PARAGRAPH)
            highlight_style.font.color.rgb = RGBColor(0, 102, 204)
            highlight_style.font.size = Pt(11)
            highlight_style.font.italic = True

    def _create_attribute(self, element: OxmlElement, name: str, value: str):
        """Create a custom XML attribute"""
        element.set(qn(name), value)

    def _setup_page_layout(self):
        """Setup page layout including headers and footers"""
        # Get the section properties
        section = self.doc.sections[0]

        # Set page margins
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(1)
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)

        # Setup header
        header = section.header
        header_para = header.paragraphs[0]
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Add logo to header (left side)
        header_table = header.add_table(1, 2, width=Inches(4))
        header_table.autofit = False
        header_table.columns[0].height=Inches(2)
        header_table.columns[0].width = Inches(3.75)
        header_table.columns[1].width = Inches(3.75)

        # Left cell for logo placeholder
        logo_cell = header_table.cell(0, 0)
        logo_para = logo_cell.paragraphs[0]
        logo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        logo_run = logo_para.add_run()
        logo_run.font.size = Pt(8)
        logo_run.text = "[Logo]"  # Placeholder for logo

        # Right cell for header text
        text_cell = header_table.cell(0, 1)
        text_para = text_cell.paragraphs[0]
        text_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        text_run = text_para.add_run()
        text_run.font.size = Pt(8)
        text_run.text = f"{self.title} {self.year}"

        # Setup footer
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add footer table
        footer_table = footer.add_table(1, 3, width=Inches(8))
        footer_table.autofit = False

        # Set column widths
        for i, width in enumerate([2.5, 2.5, 2.5]):
            footer_table.columns[i].width = Inches(width)

        # Left cell: Date
        date_cell = footer_table.cell(0, 0)
        date_para = date_cell.paragraphs[0]
        date_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        date_run = date_para.add_run()
        date_run.font.size = Pt(8)
        date_run.text = datetime.now().strftime("%d.%m.%Y")

        # Middle cell: Document info
        info_cell = footer_table.cell(0, 1)
        info_para = info_cell.paragraphs[0]
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_run = info_para.add_run()
        info_run.font.size = Pt(8)
        info_run.text = "DMU Consult"

        # Right cell: Page numbers
        page_cell = footer_table.cell(0, 2)
        page_para = page_cell.paragraphs[0]
        page_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        page_run = page_para.add_run()
        page_run.font.size = Pt(8)

        # Add page numbers
        self._add_page_number(page_para)

    def _add_page_number(self, paragraph):
        """Add page numbers to the document"""
        page_num_run = paragraph.add_run()
        fldChar1 = self._create_element('w:fldChar')
        self._create_attribute(fldChar1, 'w:fldCharType', 'begin')

        instrText = self._create_element('w:instrText')
        instrText.text = "PAGE"

        fldChar2 = self._create_element('w:fldChar')
        self._create_attribute(fldChar2, 'w:fldCharType', 'end')

        page_num_run._r.append(fldChar1)
        page_num_run._r.append(instrText)
        page_num_run._r.append(fldChar2)

        # Add "of" text
        paragraph.add_run(" von ")

        # Add total pages
        total_run = paragraph.add_run()
        fldChar3 = self._create_element('w:fldChar')
        self._create_attribute(fldChar3, 'w:fldCharType', 'begin')

        instrText2 = self._create_element('w:instrText')
        instrText2.text = "NUMPAGES"

        fldChar4 = self._create_element('w:fldChar')
        self._create_attribute(fldChar4, 'w:fldCharType', 'end')

        total_run._r.append(fldChar3)
        total_run._r.append(instrText2)
        total_run._r.append(fldChar4)

    def update_header_logo(self, logo_path: str):
        """Update the header logo if a path is provided"""
        if logo_path and os.path.exists(logo_path):
            section = self.doc.sections[0]
            header_table = section.header.tables[0]
            logo_cell = header_table.cell(0, 0)
            logo_cell.paragraphs[0].clear()
            logo_run = logo_cell.paragraphs[0].add_run()
            logo_run.add_picture(logo_path, width=Inches(1))

    def add_table_of_contents(self):
        """Generates table of contents with correct styling"""
        self.doc.add_heading("Inhaltsverzeichnis", level=1)

        # Updated section numbering
        toc_sections = [
            "1. Vorbemerkung",
            "2. Organisation",
            "3. Aktuelles 2023",
            "4. Fragen und Antworten",
            "5. Relevante Dokumente",
            "6. Termine",
            "7. Anhang"
        ]

        # Add TOC entries with custom style
        for section in toc_sections:
            p = self.doc.add_paragraph(style='Contents')
            p.add_run(section)
            p.paragraph_format.left_indent = Inches(0.5)

        self.doc.add_page_break()

    def add_title_page(self, organization: str, department: str, address: List[str], logo_path: str = None):
        """Creates an enhanced title page with logo"""
        section = self.doc.sections[0]
        section.top_margin = Cm(1.27)  # 0.5 inch / 1.27 cm top margin

        # Add logo with proper spacing
        if logo_path and os.path.exists(logo_path):
            logo_paragraph = self.doc.add_paragraph()
            logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            logo_paragraph.space_before = Pt(12)  # Small space before logo
            logo_paragraph.space_after = Pt(24)  # Small space after logo
            logo_run = logo_paragraph.add_run()
            # Add a paragraph with 2 cm space after the header
            spacer_paragraph = self.doc.add_paragraph()
            logo_run.add_picture(logo_path, width=Inches(3))

        # Add spacing after logo
        spacing_paragraph = self.doc.add_paragraph()
        spacing_paragraph.space_before = Pt(6)

        self.doc.add_paragraph().add_run().add_break()

        # Organization details
        org_para = self.doc.add_paragraph()
        org_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        org_run = org_para.add_run(organization)
        org_run.font.size = Pt(16)
        org_run.font.bold = True
        org_run.font.color.rgb = RGBColor(0, 51, 102)

        # Department and address
        for line in [department] + address:
            para = self.doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.add_run(line)

        self.doc.add_paragraph().add_run().add_break()

        # Title
        title_para = self.doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        title_run = title_para.add_run(f"{self.title} {self.year}")
        title_run.bold = True
        title_run.font.size = Pt(24)
        title_run.font.color.rgb = RGBColor(0, 51, 102)

        self.doc.add_page_break()

    def add_introduction_with_image(self, image_path: str = None, image_caption: str = ""):
        """Adds introduction section with optional image"""
        self.doc.add_heading("1. Vorbemerkung", level=1)

        intro_text = "Dieser Jahresbericht bietet einen umfassenden Überblick über die wichtigsten Entwicklungen und Aktivitäten des vergangenen Jahres. Er wurde mithilfe modernster KI-gestützter Analysemethoden erstellt, um präzise und relevante Informationen zu liefern."
        self.doc.add_paragraph(intro_text)

      #  if image_path and os.path.exists(image_path):
          #  img_paragraph = self.doc.add_paragraph()
           # img_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
           # img_run = img_paragraph.add_run()
           # img_run.add_picture(image_path)
    #
            #if image_caption:
               # caption_para = self.doc.add_paragraph()
                #caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                #caption_para.add_run(image_caption).italic = True

        #self.doc.add_paragraph().add_run().add_break()



    def add_rag_content(self, inputs: List[str], answers: List[str],
                        relevant_docs: List[str], all_docs: List[str]):
        """Adds enhanced RAG Q&A content"""
        # Q&A Section
        self.doc.add_heading("2. Organisation", level=1)
        self.doc.add_paragraph("Organisatorische Struktur und Aufbau...")

        # Current developments
        self.doc.add_heading("3. Aktuelles 2023", level=1)
        self.doc.add_paragraph("Aktuelle Entwicklungen und Fortschritte...")


        self.doc.add_heading("4. Fragen und Antworten", level=1)

        # Add Q&A pairs with enhanced formatting
        self.doc.add_heading("4.1 Fragenkatalog", level=2)
        for i, (question, answer) in enumerate(zip(inputs, answers), 1):
            # Question box
            q_para = self.doc.add_paragraph()
            q_para.style = 'Highlight'
            q_run = q_para.add_run(f"Frage {i}: ")
            q_run.bold = True
            q_para.add_run(question)

            # Answer with indentation
            a_para = self.doc.add_paragraph()
            a_para.paragraph_format.left_indent = Inches(0.25)
            a_run = a_para.add_run(f"Antwort {i}: ")
            a_run.bold = True
            a_para.add_run(answer)

            self.doc.add_paragraph()  # Spacing

        # Documents Section with enhanced formatting
        self.doc.add_heading("5. Relevante Dokumente", level=1)

        # Most relevant documents
        self.doc.add_heading("5.1 Hauptreferenzen", level=2)
        for i, doc in enumerate(relevant_docs, 1):
            p = self.doc.add_paragraph(style='List Number')
            p.add_run(doc)

        # All associated documents
        self.doc.add_heading("5.2 Weitere Referenzen", level=2)
        for i, doc in enumerate(all_docs, 1):
            p = self.doc.add_paragraph(style='List Number')
            p.add_run(str(doc))

        self.doc.add_heading("6. Termine", level=1)

        table = self.doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        header_cells = table.rows[0].cells
        headers = ['Datum', 'Veranstaltung', 'Ort']
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].font.bold = True

        # Appendix
        self.doc.add_heading("7. Anhang", level=1)
        self.doc.add_paragraph("Zusätzliche Dokumente und Materialien...")


    def save(self, filename: str = None) -> Optional[BytesIO]:
        """Saves the enhanced document"""
        if filename:
            self.doc.save(filename)
            return None
        else:
            byte_io = BytesIO()
            self.doc.save(byte_io)
            byte_io.seek(0)
            return byte_io

    def create_document(self, org_data: dict, logo_path: str = None, intro_image_path: str = None):
        """Create the complete document with all sections"""
        # Add title page
        self.add_title_page(
            org_data['name'],
            org_data['department'],
            org_data['address'],
            logo_path
        )

        # Update header logo if provided
        if logo_path:
            self.update_header_logo(logo_path)

        # Add rest of the document
        self.add_table_of_contents()
        self.add_introduction_with_image(intro_image_path)

        return self.doc


def create_rag_report(
        inputs: List[str],
        answers: List[str],
        documents: List[str],
        all_docs: List[str],
        org_data: Dict[str, Any] = None,
        logo_path: str = None,
        intro_image_path: str = None
) -> BytesIO:
    """Create complete report with headers and footers"""
    if org_data is None:
        org_data = {
            'name': 'Landeshauptstadt München',
            'department': 'Kommunalreferat IS-SP',
            'address': ['Denisstraße 2', '80335 München']
        }

    report = RAGReportGenerator()

    # Create complete document
    doc = report.create_document(org_data, logo_path, intro_image_path)

    # Add RAG content
    report.add_rag_content(inputs, answers, documents, all_docs)

    # Save to BytesIO
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio
# Example usage
if __name__ == "__main__":
    # Sample data
    sample_inputs = ["What is the total waste disposed in 2023?"]
    sample_answers = ["In 2023, approximately 21,000 tons of waste were disposed."]
    sample_docs = ["Waste Management Report 2023"]
    sample_all_docs = ["Full Environmental Report 2023", "Waste Statistics 2023"]

    # Generate report
    report_bytes = create_rag_report(
        sample_inputs,
        sample_answers,
        sample_docs,
        sample_all_docs,
        logo_path='../static/dmu.img',
        intro_image_path='../static/dmu.img'
    )

    # Save to file
    with open('rag_report_2023.docx', 'wb') as f:
        f.write(report_bytes.getvalue())
